#!/usr/bin/env python3
# -*- coding:utf-8 -*-
# pip install python-docx
# pip install openpyxl
from openpyxl import load_workbook
from openpyxl import Workbook  # 导入excel包
from docx import Document  # 导入Docx包
from docx.shared import Cm, Inches, Pt  # 导入单位换算函数
from docx.oxml.ns import qn  # docx中文字体模块
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入对齐选项
import time  # 导入时间

try:
    # 路径
    LoadPath = "订单信息.xlsx"  # 加载excel路径(这里为相对路径，excel表与该程序在同一文件夹下就能识别，所以只用excel文件名即可)
    SavePath = "ResultTest.xlsx"  # 另存为路径（用于保存修改过的excel）
    WordSavePath = '订单周报.docx'  # word另存为路径（自动生成Word的路径）
    # excel表格初始化
    book = load_workbook(LoadPath)  # 加载已有Excel文档
    try:
        sheet = book['订单信息']  # 加载需要的工作簿（这里为excel表中的sheet工作簿）
    except:
        print('优先处理异常：\nExcel中工作簿(sheet)名称必须为：订单信息。请修改Excel后重新运行程序')
    print("————————————————————————————————————————————————————————\n抓取工作簿名称：", sheet.title)  # sheet.title为工作簿名称

    # （1）剔除首行：删除第一行不需要的数据
    sheet.delete_rows(1)  # 删除行，（）里面数据对应第几行，这里为第一行
    print('执行剔除首行成功')  # 用来在程序中表现执行成功
    # book.save(SavePath)

    print('订单数量：', sheet.max_row)  # 程序输出工作簿总共有几行
    print('最大列：', sheet.max_column)  # 程序输出工作簿总共有几列

    # （2）遍历列,更新列值。统一‘订单归属州’的参数
    # 大理白族自治州—>大理州
    # 迪庆藏族自治州—>迪庆州
    # 楚雄彝族自治州—>楚雄州
    # for cell in sheet['D']:  # D列为“订单归属州”
    #     if cell.value == '大理白族自治州':
    #         cell.value = '大理州'
    #     elif cell.value == '迪庆藏族自治州':
    #         cell.value = '迪庆州'
    #     elif cell.value == '楚雄彝族自治州':
    #         cell.value = '楚雄州'
    # print('执行统一“地州”成功')
    # book.save(SavePath)

    # （3）获取表中所有归属地-地州列表：去重复
    formatList = []  # 中间list
    for col in sheet['D']:  # D列为“订单归属州”
        formatList.append(col.value)
    GsPalceList = list(set(formatList))  # 使用set的特型，python的set和其他语言类似, 是一个无序不重复元素集。将list转换为set再转换回来完成。
    print('抓取归属地州清单：\n', GsPalceList)

    # （4）提取数据日期
    DateList = []
    for cell in sheet['S']:  # S列为“发起订单时间”
        DateList.append(cell.value)

    MinDate = min(DateList)  # 获取订单时间最小日期
    MinDateYear = MinDate[0:4]  # 获取最小日期年份
    MinDateMonth = MinDate[5:7]  # 获取最小日期月份
    MinDateDay = MinDate[8:10]  # 获取最小日期日份

    MaxDate = max(DateList)  # 获取订单时间最大日期
    MaxDateYear = MaxDate[0:4]  # 获取最大日期年份
    MaxDateMonth = MaxDate[5:7]  # 获取最大日期月份
    MaxDateDay = MaxDate[8:10]  # 获取最大日期日份
    print(
        f'抓取日期范围：{MinDateYear}年{MinDateMonth}月{MinDateDay}日-{MaxDateMonth}月{MaxDateDay}日')  # 以 f 开头，包含的{}表达式在程序运行时会被表达式的值代替。

    # （5）涉及订单金额计算
    SumGoodsMoney = 0  # 订单金额之和
    for col in sheet['L']:  # L列为“订单金额”
        if (col.value != ""):  # 排除xlsx单元格内空值
            m = float(col.value)
            SumGoodsMoney = SumGoodsMoney + m

    print(f"涉及订单金额和：{SumGoodsMoney}")

    # 成功订单金额
    SumGoodsRefund = 0  # 成功收款金额之和
    for col in sheet['W']:  # W列为“用券折扣后价格”
        if (col.value != ""):  # 排除xlsx单元格内空值
            n = float(col.value)
            SumGoodsRefund = SumGoodsRefund + n
    print(f"成功收款金额和：{SumGoodsRefund}")

    # （6）统计订单结果
    SumResultInHand = 0  # 订单结果：发货中
    SumResultBackout = 0  # 订单结果：订单撤销
    SumResultSuccess = 0  # 订单结果：订单完成
    SumResultRefuse = 0  # 订单结果:订单缺货
    SumResultWPending = 0  # 订单结果：待付款
    for col in sheet['V']:  # v列为订单结果列
        if col.value == '发货中':
            SumResultInHand += 1
        elif col.value == '订单撤销':
            SumResultBackout += 1
        elif col.value == '订单完成':
            SumResultSuccess += 1
        elif col.value == '订单缺货':
            SumResultRefuse += 1
        elif col.value == '待付款':
            SumResultWPending += 1

    # 订单来源统计：（美团在线订单、电话订单、小程序订单）
    AppOnlineOrder = 0  # 水果APP在线订单
    WeChetOrder = 0  # 小程序订单
    PhoneOrder = 0  # 电话订单
    AppVoice = 0  # 现场订单
    SystemIn = 0  # 系统录入
    for col in sheet['G']:
        if col.value == '水果APP在线订单':
            AppOnlineOrder += 1
        elif col.value == '小程序在线订单':
            WeChetOrder += 1
        elif col.value == '电话订单':
            PhoneOrder += 1
        elif col.value == '现场订单':
            AppVoice += 1
        elif col.value == '系统录入':
            SystemIn += 1
    # print('水果APP在线订单：{0},小程序在线订单:{1},电话订单:{2},现场订单:{3}'.format(AppOnlineOrder,WeChetOrder,PhoneOrder,AppVoice))

    # *****************************************************************
    # 保存输出成为.Docx文件

    # 将内容格式化输出到文本中
    WordSavePath = (f'水果订单销售报告_{MinDateYear}_{MinDateMonth}月{MinDateDay}日-{MaxDateMonth}月{MaxDateDay}日.docx')
    Word = Document()

    # 全局设置字体
    Word.styles['Normal'].font.name = u'宋体'
    Word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 设置A4纸张
    section = Word.sections[0]
    section.page_width = Cm(21)  # 设置A4纸的宽度
    section.page_height = Cm(29.7)  # 设置A4纸的高度
    print('导出Word页面的宽度和高度（A4）：', section.page_width.cm, section.page_height.cm)

    # 首段
    str1 = Word.add_paragraph(style=None)  # 增加一个段落
    str1_run = str1.add_run('云南省水果商城销售报告')  # 增加文字块
    str1_run.bold = True  # 加粗
    str1_run.font.size = Pt(18)  # 行距
    str1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    str2 = Word.add_paragraph(style=None)
    str2_run = str2.add_run(f'{MinDateYear}年{MinDateMonth}月{MinDateDay}日-{MaxDateMonth}月{MaxDateDay}日\n')
    str2_run.bold = True
    str2_run.font.size = Pt(14)
    str2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    str3 = Word.add_paragraph()
    str3.paragraph_format.line_spacing = Pt(28)  # 行距
    str3.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
    if SumGoodsRefund < 10000:
        str3_run = str3.add_run(
            f'{MinDateYear}年{MinDateMonth}月{MinDateDay}日-{MaxDateMonth}月{MaxDateDay}日 24:00，水果商城平台共收到订单{sheet.max_row}起，涉及金额{"{:.2f}万元".format(SumGoodsMoney / 10000)}，订单完成收款金额{"{:.2f}元".format(SumGoodsRefund)}，其中发货订单{SumResultInHand + SumResultSuccess}起（发货中{SumResultInHand}起，订单完成{SumResultSuccess}起），待受理{SumResultWPending}起，订单撤销{SumResultBackout}起，订单缺货{SumResultRefuse}起。')
    else:
        str3_run = str3.add_run(
            f'{MinDateYear}年{MinDateMonth}月{MinDateDay}日-{MaxDateMonth}月{MaxDateDay}日 24:00，水果商城平台共收到订单{sheet.max_row}起，涉及金额{"{:.2f}万元".format(SumGoodsMoney / 10000)}，订单完成收款金额{"{:.2f}万元".format(SumGoodsRefund / 10000)}，其中发货订单{SumResultInHand + SumResultSuccess}起（发货中{SumResultInHand}起，订单完成{SumResultSuccess}起），待受理{SumResultWPending}起，订单撤销{SumResultBackout}起，订单缺货{SumResultRefuse}起。')
    str3_run.font.name = (u'仿宋')  # 字体样式
    str3_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
    str3_run.font.size = Pt(14)  # 字体大小

    # 第一部分、按订单来源类型划分
    str4 = Word.add_paragraph()
    str4.paragraph_format.line_spacing = Pt(28)  # 行距
    str4.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
    # 一、订单种类分布
    str4_run = str4.add_run('一、水果订单种类分布：')
    str4_run.font.name = (u'黑体')  # 字体样式
    str4_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 字体样式
    str4_run.font.size = Pt(14)  # 字体大小
    str4_run.font.bold = True

    str5 = Word.add_paragraph()
    str5.paragraph_format.line_spacing = Pt(28)  # 行距
    str5.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
    str5_run = str5.add_run(
        f'“水果商城”平台“水果APP在线订单”{AppOnlineOrder}起，现场订单{AppVoice}起，电话订单{PhoneOrder}起，小程序订单{WeChetOrder}起，系统录入{SystemIn}起。')
    str5_run.font.name = (u'仿宋')  # 字体样式
    str5_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
    str5_run.font.size = Pt(14)  # 字体大小

    # 第二部分、按订单地区划分统计：序号+地区+金额
    str6 = Word.add_paragraph()
    str6.paragraph_format.line_spacing = Pt(28)  # 行距
    str6.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
    str6_run = str6.add_run('二、水果订单地区分布')
    str6_run.font.name = (u'黑体')  # 字体样式
    str6_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 字体样式
    str6_run.font.size = Pt(14)  # 字体大小
    str6_run.font.bold = True
    # 按地州进行遍历统计订单金额
    for GsPalce in GsPalceList:  # 按前面获得的不重复地州列表进行遍历
        GsPalceSum = 0  # 按州市金额总数初始化为：0
        GsPalceCount = 0  # 按州市订单件数初始化为：0
        for gsd in sheet['D']:  # 遍历归属地州市，D列为归属地州市
            if GsPalce == gsd.value:  # 筛选当前归属地
                GsPalceCount += 1  # 当前归属地订单件数量
                if sheet.cell(row=gsd.row, column=(sheet['L1'].column)).value != "":  # 通过归属地单元格，获取归属地金额值，并排除空值
                    m = float(sheet.cell(row=gsd.row, column=(sheet['L1'].column)).value)  # 单元格数据转换为浮点型，L列为金额列
                    GsPalceSum += m  # 该单元格计入当前归属地总金额
        str7 = Word.add_paragraph()
        str7.paragraph_format.line_spacing = Pt(28)  # 行距
        str7.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进

        # 序号+归属地
        str7_run = str7.add_run(f'{GsPalceList.index(GsPalce) + 1}.{GsPalce}{GsPalceCount}件')  # 序号+归属地：进行加粗设置
        str7_run.font.name = (u'仿宋')  # 字体样式
        str7_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
        str7_run.font.size = Pt(14)  # 字体大小
        str7_run.bold = True  # 加粗首段州市文字
        # 金额
        str7_run = str7.add_run(f'：{"{:.2f}元".format(GsPalceSum)}')  # 州市统计的金额
        str7_run.font.name = (u'仿宋')  # 字体样式
        str7_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
        str7_run.font.size = Pt(14)  # 字体大小

    # 第三部分、按类型划分
    TypeList = ['酸性水果', '亚酸性水果', '甜性水果', '温热性水果', '凉性水果']  # 订单类型列表

    # 三、水果类型分布
    str8 = Word.add_paragraph()
    str8.paragraph_format.line_spacing = Pt(28)  # 行距
    str8.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
    str8_run = str8.add_run('三、水果类型分布')
    str8_run.font.name = (u'黑体')  # 字体样式
    str8_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 字体样式
    str8_run.font.size = Pt(14)  # 字体大小
    str8_run.font.bold = True

    # 根据类型的种类按地州进行小节统计
    for Type in TypeList:
        TypeSum = 0  # 类型总数量
        TypeMoneySum = 0  # 类型总金额
        for cell in sheet['k']:  # k列为商品类型列
            if Type == cell.value:
                TypeSum += 1
                if sheet.cell(row=cell.row, column=(sheet['L1'].column)).value != "":  # 通过该类型单元格，并取该行金额值
                    m = float(sheet.cell(row=cell.row, column=(sheet['L1'].column)).value)  # 单元格数据转换为浮点型，L列为金额列
                    TypeMoneySum += m  # 该单元格计入当前类型总金额
        str9 = Word.add_paragraph()
        str9.paragraph_format.line_spacing = Pt(28)  # 行距
        str9.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
        str9_run = str9.add_run(
            f'{TypeList.index(Type) + 1}.{Type}{TypeSum}起，涉及金额{"{:.2f}元".format(TypeMoneySum)}：')  # 序号+类型+数量+金额
        str9_run.font.name = (u'仿宋')  # 字体样式
        str9_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
        str9_run.font.size = Pt(14)  # 字体大小
        str9_run.bold = True  # 加粗金额类型统计字段
        for TypeGsPalce in GsPalceList:
            TypeGsPalceSum = 0  # 当前类型归属地按州市金额总数初始化为：0
            TypeGsPalceCount = 0  # 当前类型按州市订单件数
            TypeumResultInHand = 0  # 当前类型订单结果：发货中
            TypeSumResultBackout = 0  # 当前类型订单结果：订单撤销
            TypeSumResultSuccess = 0  # 当前类型订单结果：订单成功
            TypeSumResultRefuse = 0  # 当前类型订单结果:订单缺货
            TypeSumResultWPending = 0  # 当前类型订单结果：待付款
            for cell in sheet['D']:  # 遍历归属地州市
                if TypeGsPalce == cell.value and sheet.cell(row=cell.row, column=(
                        sheet['k1']).column).value == Type:  # 筛选当前归属地,并且该行商品类型为当前类型
                    TypeGsPalceCount += 1  # 当前归属地订单件数量
                    if sheet.cell(row=cell.row, column=(sheet['L1'].column)).value != "":  # 通过归属地单元格，获取归属地金额值，并排除空值
                        m = float(sheet.cell(row=cell.row, column=(sheet['L1'].column)).value)  # 单元格数据转换为浮点型，L列为金额列
                        TypeGsPalceSum += m  # 该单元格计入当前归属地总金额
                    if sheet.cell(row=cell.row, column=(sheet['V1'].column)).value == "发货中":
                        TypeumResultInHand += 1
                    elif sheet.cell(row=cell.row, column=(sheet['V1'].column)).value == "订单撤销":
                        TypeSumResultBackout += 1
                    elif sheet.cell(row=cell.row, column=(sheet['V1'].column)).value == "订单完成":
                        TypeSumResultSuccess += 1
                    elif sheet.cell(row=cell.row, column=(sheet['V1'].column)).value == "订单缺货":
                        TypeSumResultRefuse += 1
                    elif sheet.cell(row=cell.row, column=(sheet['V1'].column)).value == "待付款":
                        TypeSumResultWPending += 1
            if TypeGsPalceCount == 0:
                continue
            else:
                str10 = Word.add_paragraph()
                str10.paragraph_format.line_spacing = Pt(28)  # 行距
                str10.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进
                str10_run = str10.add_run(f'{TypeGsPalce}')  # 根据类型的种类按地州进行小节统计
                str10_run.font.name = (u'仿宋')  # 字体样式
                str10_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
                str10_run.font.size = Pt(14)  # 字体大小
                str10_run.font.bold = True
                str10_run = str10.add_run(
                    f'{Type}{TypeGsPalceCount}起，涉及金额{"{:.2f}元".format(TypeGsPalceSum)}，发货中{TypeumResultInHand}起，待付款{TypeSumResultWPending}起，订单撤销{TypeSumResultBackout}起，订单成功{TypeSumResultSuccess}起，订单缺货{TypeSumResultRefuse}起。')  # 根据类型的种类按地州进行小节统计
                str10_run.font.name = (u'仿宋')  # 字体样式
                str10_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 字体样式
                str10_run.font.size = Pt(14)  # 字体大小
    Word.save(WordSavePath)
except:
    print('\n注意！\n确保系统导出的Excel文件已经放到与该程序同目录下，重新运行该程序（文件必须为：订单信息.xlsx）\nauthor:贤仔 e-mail：972912623@qq.com')
    # 显示倒计时
    for i in range(0, 10):
        print(f'\r系统{10 - i}秒后自动退出。', end="")
        time.sleep(1)

else:
    print(
        '————————————————————————————————————————————————————————\n\n\n执行成功，请查看该目录下：订单数据报告.docx')
    for i in range(0, 10):
        print(f'\r系统{10 - i}秒后自动退出。', end="")
        time.sleep(1)
