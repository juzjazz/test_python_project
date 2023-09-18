import os

from openpyxl import load_workbook
from openpyxl import Workbook  # 导入excel包
from docx import Document  # 导入Docx包
from docx.shared import Cm, Inches, Pt  # 导入单位换算函数
from docx.oxml.ns import qn  # docx中文字体模块
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入对齐选项
import time  # 导入时间


def creat_word(row_list):
    for row in row_list:
        word_save_path = (f'{row[5].value}--江苏公司.docx')
        if not os.path.exists(word_save_path):
            word = Document()
            # 全局设置字体
            word.styles['Normal'].font.name = u'宋体'
            word.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            # 设置A4纸张
            section = word.sections[0]
            section.page_width = Cm(21)  # 设置A4纸的宽度
            section.page_height = Cm(29.7)  # 设置A4纸的高度
            print('导出word页面的宽度和高度（A4）：', section.page_width.cm, section.page_height.cm)
            print(row[6].value)

            # 首段
            str1 = word.add_paragraph(style=None)  # 增加一个段落
            str1_run = str1.add_run(f'{row[5].value}')  # 增加文字块
            str1_run.bold = True  # 加粗
            str1_run.font.size = Pt(18)  # 行距
            str1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

            str2 = word.add_paragraph(style=None)
            str2_run = str2.add_run(f'一、能力名称及编码\n')
            str2_run.bold = True
            str2_run.font.size = Pt(14)

            str3 = word.add_paragraph(style=None)
            str3_run = str3.add_run(f'{row[5].value}{row[6].value}\n')
            str3.paragraph_format.line_spacing = Pt(28)  # 行距
            str3.paragraph_format.first_line_indent = Cm(1.10)  # 首行缩进

            str4 = word.add_paragraph(style=None)
            str4_run = str4.add_run(f'二、能力基本情况\n')
            str4_run.bold = True
            str4_run.font.size = Pt(14)

            # str5 = word.add_paragraph(style=None)
            # str5_run = str5.add_run(f'能力xxxxxxxxxxxxxxxxxxx\n')
            # str5_run.bold = True
            # str5_run.font.size = Pt(14)

            str6 = word.add_paragraph(style=None)
            str6_run = str6.add_run(f'三、本次申报能力贡献度说明\n')
            str6_run.bold = True
            str6_run.font.size = Pt(14)

            # 建表table
            table = word.add_table(rows=4, cols=len(row_list) + 1, style='Table Grid')
            # 表格各个行的标题
            field_list = ["合同编码", "能力贡献度", "能力收入（不含税）", "能力收入（含税）"]
            r = table.columns[0].cells
            for i in range(len(field_list)):
                r[i].text = field_list[i]
            # 表格中插入每个合同内容
            for i, v in enumerate(row_list):
                temp = table.columns[i + 1].cells
                j = 0
                temp[j].text = str(v[1].value)
                temp[j + 1].text = f'{v[11].value * 100}%'
                temp[j + 2].text = str(v[10].value)
                temp[j + 3].text = str(v[9].value)

        word.save(word_save_path)


try:
    # 能力编码数据
    id = ['A225000055', 'A325000324']

    LoadPath = "test.xlsx"  # 加载excel路径(这里为相对路径，excel表与该程序在同一文件夹下就能识别，所以只用excel文件名即可)
    # excel表格初始化
    workbook = load_workbook(LoadPath)  # 加载已有Excel文档
    sheet = workbook['Sheet1']
    print('订单数量：', sheet.max_row)  # 程序输出工作簿总共有几行
    print('最大列：', sheet.max_column)  # 程序输出工作簿总共有几列
    # 遍历每一行的能力数据编码，并且匹配是否符合要求，一个能力编码对应一个word
    for i in id:
        row_list = []
        for row in sheet.rows:
            if i == row[6].value:
                row_list.append(row)
        creat_word(row_list)


except():
    print('\n注意！\n确保系统导出的Excel文件已经放到与该程序同目录下，重新运行该程序')
